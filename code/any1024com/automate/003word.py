# coding=utf-8
import pathlib
from docx import Document


def gen_invite_letter(template_path, out_path):
    """
    自动批量生成邀请函

    :params template_path: 模板文件地址
    :params out_path: 输出文件目录地址

    Version: 0.0.1
    Author : yichu.cheng
    """
    invited_comps = ['广州只差一个程序员了有限公司', '广州一初有限公司', '广州学用有限公司']
    data_dict = {
        'company': '广东哈喽沃得公司',
        'product': 'Python1024',
        'year': 2020,
        'month': 10,
        'day': 24,
        'time': 10,
        'email': 'ichengplus@qq.com',
        'contact': '程一初',
        'phone': '88881024'
    }

    for invited in invited_comps:
        doc = Document(template_path)
        paragraphs = doc.paragraphs
        file_out_path = out_path.joinpath(f'003word_case_out_{invited}.docx')
        for p in paragraphs:
            text = p.text
            new_text = text.format(invited=invited, **data_dict)
            p.clear()
            p.add_run(new_text)

        doc.save(file_out_path)


if __name__ == "__main__":
    path = pathlib.Path().cwd().joinpath('data/automate/003word')
    template_path = path.joinpath('003word_case_template.docx')
    out_path = path.joinpath('003word_case_out')
    if not out_path.is_dir():
        out_path.mkdir()
    gen_invite_letter(template_path, out_path)
